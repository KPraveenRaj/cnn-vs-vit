"""Shared Word helpers: a consistent academic report skeleton in one place.

Same discipline as deckkit — geometry and styling live here so the report
builders read as content. Heading numbering is explicit rather than relying on
Word list styles, which do not survive round-tripping between editors.
"""
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

ACCENT = RGBColor(0x1F, 0x4E, 0x79)
BODY = RGBColor(0x1A, 0x1A, 0x1A)
MUTED = RGBColor(0x59, 0x59, 0x59)


def new_doc():
    doc = Document()
    for s in doc.sections:
        s.top_margin = s.bottom_margin = Inches(1.0)
        s.left_margin = s.right_margin = Inches(1.1)
    st = doc.styles["Normal"]
    st.font.name = "Calibri"
    st.font.size = Pt(11)
    st.paragraph_format.space_after = Pt(7)
    st.paragraph_format.line_spacing = 1.15
    return doc


def title_page(doc, title, subtitle, meta_lines):
    for _ in range(4):
        doc.add_paragraph()
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(title); r.font.size, r.font.bold, r.font.color.rgb = Pt(22), True, ACCENT
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(subtitle); r.font.size, r.font.color.rgb = Pt(14), BODY
    doc.add_paragraph()
    for line in meta_lines:
        p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(line); r.font.size, r.font.color.rgb = Pt(11.5), MUTED
    doc.add_page_break()


def h1(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(16)
    r = p.add_run(text)
    r.font.size, r.font.bold, r.font.color.rgb = Pt(16), True, ACCENT
    return p


def h2(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(11)
    r = p.add_run(text)
    r.font.size, r.font.bold, r.font.color.rgb = Pt(12.5), True, BODY
    return p


def para(doc, text, italic=False, size=11):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.font.size, r.font.italic, r.font.color.rgb = Pt(size), italic, BODY
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    return p


def bullet(doc, text, level=0):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.left_indent = Inches(0.3 + 0.28 * level)
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run(text); r.font.size, r.font.color.rgb = Pt(11), BODY
    return p


def figure(doc, path, caption_text, width_in=6.1):
    if not path:
        para(doc, f"[figure pending: {caption_text}]", italic=True)
        return
    doc.add_picture(str(path), width=Inches(width_in))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(caption_text)
    r.font.size, r.font.italic, r.font.color.rgb = Pt(9.5), True, MUTED


def table(doc, headers, rows, caption_text=None, widths=None):
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = "Light Grid Accent 1"
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(headers):
        c = t.rows[0].cells[i]
        c.text = ""
        r = c.paragraphs[0].add_run(str(h))
        r.font.bold, r.font.size = True, Pt(10)
    for row in rows:
        cells = t.add_row().cells
        for i, v in enumerate(row):
            cells[i].text = ""
            r = cells[i].paragraphs[0].add_run(str(v))
            r.font.size = Pt(10)
    if widths:
        for i, w in enumerate(widths):
            for row in t.rows:
                row.cells[i].width = Inches(w)
    if caption_text:
        p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(caption_text)
        r.font.size, r.font.italic, r.font.color.rgb = Pt(9.5), True, MUTED
    doc.add_paragraph()
    return t


def save(doc, path):
    from pathlib import Path
    Path(path).parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(path))
    print(f"  {Path(path).name}")
